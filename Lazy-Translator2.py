import sys
import os
import time
from PyPDF2 import PdfReader
from nltk.tokenize import sent_tokenize
import openai
from docx import Document
from PyQt5.QtWidgets import QApplication, QWidget, QPushButton, QLabel, QFileDialog, QVBoxLayout
from PyQt5.QtGui import QIcon

# 設置 OpenAI API 金鑰
openai.api_key = "API KEY"

def get_escaped_path(raw_path):
    """
    從原始路徑中獲取轉義後的檔案路徑。
    """
    escaped_path = r"{}".format(raw_path)
    return escaped_path

def translate_pdf_to_chinese(pdf_path):
    # 讀取 PDF 檔案
    with open(pdf_path, 'rb') as f:
        reader = PdfReader(f)
        number_of_pages = len(reader.pages)
        chunks = []

        for i in range(number_of_pages):
            page = reader.pages[i]
            text = page.extract_text()
            sentences = sent_tokenize(text)
            input_sentences = ''

            for sentence in sentences:
                input_sentences += sentence
                if len(input_sentences) > 1000:
                    chunks.append(input_sentences)
                    input_sentences = ''
            chunks.append(input_sentences)

        # 使用 OpenAI API 完成翻譯
        translated_text = ""
        for chunk in chunks:
            prompt = "翻譯以下文字為中文：\n\n" + chunk
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "翻譯以下文字為中文："},
                    {"role": "user", "content": chunk}
                ],
                temperature=0.7,
                max_tokens=200,
                top_p=1.0,
                frequency_penalty=0.0,
                presence_penalty=0.0,
                stop=["\n"]
            )
            translation = response.choices[0].message['content']
            translated_text += translation + "\n"

        return translated_text

def save_to_word(text, title, output_dir):
    doc = Document()
    doc.add_heading(title, level=1)
    
    # 添加文字內容
    doc.add_paragraph(text)
        
    # 保存Word檔案
    doc.save(os.path.join(output_dir, f"{title}.docx"))

class Example(QWidget):
    
    def __init__(self):
        super().__init__()

        self.initUI()

    def initUI(self):
        self.setGeometry(300, 300, 400, 150)
        self.setWindowTitle("PDF 翻譯器")
        self.setWindowIcon(QIcon(r"C:\Users\Leon\Desktop\python\icon和圖片\doge.jpg"))

        layout = QVBoxLayout()

        self.label = QLabel("選擇一個 PDF 文件進行翻譯", self)
        layout.addWidget(self.label)

        self.btn = QPushButton('選擇 PDF', self)
        self.btn.clicked.connect(self.showDialog)
        layout.addWidget(self.btn)

        self.setLayout(layout)
        self.show()

    def showDialog(self):
        pdf_path, _ = QFileDialog.getOpenFileName(self, '選擇文件', '', 'PDF files (*.pdf)')
        if pdf_path:
            start_time = time.time()
            escaped_path = get_escaped_path(pdf_path)
            translated_text = translate_pdf_to_chinese(escaped_path)
            title = os.path.basename(escaped_path).split('.')[0]
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            save_to_word(translated_text, title, desktop_path)
            end_time = time.time()
            elapsed_time = end_time - start_time
            self.label.setText(f"Word檔案保存完成，共花費 {elapsed_time:.2f} 秒")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = Example()
    sys.exit(app.exec_())
